#!/usr/bin/env python3
"""
BILDR Keretszerződés generátor.

Másolja a bázis keretszerződés DOCX template-et és kicseréli benne
a Megrendelő placeholder-eit a megadott cégadatokra.

A python-docx-et használja, mert így megőrzi a stílusokat,
a számozott listákat és a formázást.

Használat:
  python3 generate_contract.py \\
    --template /path/to/your/base-template.docx \\
    --config /tmp/contract-data.json \\
    --output /path/to/output.docx

A JSON config szerkezet:
{
  "company_name": "Példa Kft.",
  "company_long_name": "Példa Korlátolt Felelősségű Társaság",
  "address": "1051 Budapest, Példa utca 1.",
  "tax_id": "12345678-1-23",
  "company_reg": "01-09-123456",
  "representative_name": "Példa Béla",
  "representative_role": "ügyvezető",
  "contact_name": "Példa Anna",
  "contact_phone": "+36 1 234 5678",
  "contact_email": "kapcsolat@pelda.hu",
  "city": "Budapest",
  "year": "2026",
  "month": "április"
}

A Vállalkozó (BILDR HUB Kft.) adatait a template eredetileg tartalmazza.

SETUP — fontos!
A `BASE_REPLACEMENTS` dict alább a *te* bázis DOCX-edben szereplő string-eket
azonosítja. Ha új template-et használsz, frissítsd a kulcsokat (a bal oldali
oszlopot) a saját bázis DOCX-edben szereplő tényleges szövegekre. Az értékek
(jobb oldal) a JSON config kulcsait használják, ezeket nem kell módosítani.
"""

import argparse
import json
import shutil
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx nincs telepítve. Futtasd: pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


# A bázis template-ben szereplő string-ek, amiket a Megrendelő adatra cserélünk.
# Ezt a dict-et a saját bázis DOCX-edben szereplő string-ekre kell igazítani.
BASE_REPLACEMENTS = {
    "Példa Ügyfél Korlátolt Felelősségű Társaság": "{company_long_name}",
    "Példa Ügyfél Kft.": "{company_name}",
    "1051 Budapest, Példa utca 1.": "{address}",
    "12345678-2-43": "{tax_id}",
    "01-09-123456": "{company_reg}",
    "Példa Béla, ügyvezető": "{representative_name}, {representative_role}",
    "Példa Béla": "{representative_name}",
    "Példa Anna": "{contact_name}",
    "+36 1 234 5678": "{contact_phone}",
    "kapcsolat@peldaugyfel.hu": "{contact_email}",
    "Példa Ügyfél marketing weboldal fejlesztése": "{project_title}",
    "peldaugyfel.hu": "{client_domain}",
    "Budapest, 2026. április": "{city}, {year}. {month}",
}


def replace_in_paragraph(paragraph, old: str, new: str):
    """
    Run-szinten cserél, hogy a formázás megmaradjon.
    Ha az old string több run-ra esik, először összevonja a run-okat.
    """
    full_text = paragraph.text
    if old not in full_text:
        return False

    # Ha egy run-ban benne van, egyszerű csere
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True

    # Több run-ra szétoszlik — összevonjuk az első run-ba és töröljük a többit
    if old in full_text:
        first_run = paragraph.runs[0] if paragraph.runs else None
        if first_run is None:
            return False
        first_run.text = full_text.replace(old, new)
        for run in paragraph.runs[1:]:
            run.text = ""
        return True
    return False


def replace_in_doc(doc: Document, old: str, new: str):
    count = 0
    for paragraph in doc.paragraphs:
        if replace_in_paragraph(paragraph, old, new):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_in_paragraph(paragraph, old, new):
                        count += 1
    return count


def main():
    parser = argparse.ArgumentParser(description="BILDR keretszerződés generátor")
    parser.add_argument("--template", required=True, help="Bázis template DOCX path")
    parser.add_argument("--config", required=True, help="JSON config file")
    parser.add_argument("--output", required=True, help="Output DOCX path")
    args = parser.parse_args()

    template_path = Path(args.template)
    config_path = Path(args.config)
    output_path = Path(args.output)

    if not template_path.exists():
        print(f"ERROR: template not found: {template_path}", file=sys.stderr)
        sys.exit(1)
    if not config_path.exists():
        print(f"ERROR: config not found: {config_path}", file=sys.stderr)
        sys.exit(1)

    config = json.loads(config_path.read_text(encoding="utf-8"))

    # Ellenőrzés
    required_fields = [
        "company_name", "company_long_name", "address", "tax_id", "company_reg",
        "representative_name", "representative_role"
    ]
    missing = [f for f in required_fields if not config.get(f)]
    if missing:
        print(f"ERROR: missing required fields: {', '.join(missing)}", file=sys.stderr)
        sys.exit(1)

    # Default-ok
    config.setdefault("contact_name", config["representative_name"])
    config.setdefault("contact_phone", "")
    config.setdefault("contact_email", "")
    config.setdefault("project_title", "Vállalkozói tevékenység")
    config.setdefault("client_domain", "")
    config.setdefault("city", "Budapest")
    config.setdefault("year", "2026")
    config.setdefault("month", "")

    # Output dir biztos
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Másolás majd megnyitás
    shutil.copy(template_path, output_path)
    doc = Document(str(output_path))

    total = 0
    for old, placeholder in BASE_REPLACEMENTS.items():
        new = placeholder.format(**config)
        n = replace_in_doc(doc, old, new)
        total += n
        if n == 0:
            print(f"WARNING: '{old[:60]}' nem található", file=sys.stderr)

    doc.save(str(output_path))
    print(f"OK: {output_path} — {total} replacement")


if __name__ == "__main__":
    main()
