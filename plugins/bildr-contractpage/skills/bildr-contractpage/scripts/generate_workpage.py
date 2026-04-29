#!/usr/bin/env python3
"""
BILDR Munkalap (Contract Page) generátor.

Egy meglévő keretszerződéshez kapcsolódó N. sz. Munkalapot generál DOCX-ben.
A munkalap szabadon strukturált — ezért nem template-alapú, hanem
python-docx-szel építjük fel.

Használat:
  python3 generate_workpage.py --config /tmp/workpage-data.json --output /path/to/munkalap-N.docx

JSON config:
{
  "workpage_number": 2,
  "company_name": "Példa Ügyfél Kft.",
  "company_long_name": "Példa Ügyfél Korlátolt Felelősségű Társaság",
  "framework_contract_date": "2026. április 8.",
  "project_title": "Marketing weboldal fejlesztése",
  "tasks": [
    "Kliens feedback felmérés, priorizálás és tervegyeztetés",
    "Marketing landing page fejlesztése"
  ],
  "tech_stack": "Next.js 16, React 19, Tailwind CSS v4, Cloudflare Workers",
  "third_party_costs": [
    {"item": "Cloudflare Workers hosting", "estimate": "~0 Ft/hó (free tier)"}
  ],
  "fee_net": "200 000 Ft",
  "fee_text": "kettőszázezer forint",
  "payment_schedule": "A szerződés aláírásakor egy összegben: 200 000 Ft + ÁFA",
  "deadline": "2026. május 31.",
  "city": "Budapest",
  "year": "2026",
  "month": "május"
}
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx nincs telepítve. Futtasd: pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Number")
    return p


def add_signature_line(doc: Document, label: str, name: str, role: str):
    add_para(doc, "_______________________")
    add_para(doc, name, bold=True)
    add_para(doc, role)
    add_para(doc, label, italic=True)


def main():
    parser = argparse.ArgumentParser(description="BILDR munkalap generátor")
    parser.add_argument("--config", required=True, help="JSON config file")
    parser.add_argument("--output", required=True, help="Output DOCX path")
    args = parser.parse_args()

    config_path = Path(args.config)
    output_path = Path(args.output)

    if not config_path.exists():
        print(f"ERROR: config not found: {config_path}", file=sys.stderr)
        sys.exit(1)

    cfg = json.loads(config_path.read_text(encoding="utf-8"))

    required = ["workpage_number", "company_name", "company_long_name",
                "framework_contract_date", "project_title", "tasks", "fee_net"]
    missing = [f for f in required if not cfg.get(f)]
    if missing:
        print(f"ERROR: missing required fields: {', '.join(missing)}", file=sys.stderr)
        sys.exit(1)

    cfg.setdefault("tech_stack", "")
    cfg.setdefault("third_party_costs", [])
    cfg.setdefault("payment_schedule", f"A szerződés aláírásakor egy összegben: {cfg['fee_net']} + ÁFA")
    cfg.setdefault("deadline", "")
    cfg.setdefault("city", "Budapest")
    cfg.setdefault("year", "")
    cfg.setdefault("month", "")
    cfg.setdefault("fee_text", "")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Címek
    add_heading(doc, f"{cfg['workpage_number']}. számú melléklet", level=2)
    add_heading(doc, f"{cfg['workpage_number']}. sz. MUNKALAP", level=1)

    add_para(doc, "")

    # Hivatkozás a keretszerződésre
    p = doc.add_paragraph()
    p.add_run("A jelen Munkalap a ").italic = True
    r = p.add_run(cfg["framework_contract_date"])
    r.bold = True
    r.italic = True
    p.add_run(" napján kelt Vállalkozási Keretszerződéshez (a továbbiakban: „Keretszerződés”) kapcsolódik. A Keretszerződésben rögzített rendelkezések jelen Munkalapra is irányadók.").italic = True

    add_para(doc, "")

    # Felek
    add_heading(doc, "Felek", level=2)

    p = doc.add_paragraph()
    p.add_run("Megrendelő: ").bold = True
    p.add_run(cfg["company_long_name"])

    p = doc.add_paragraph()
    p.add_run("Vállalkozó: ").bold = True
    p.add_run("BILDR HUB Korlátolt Felelősségű Társaság")

    add_para(doc, "")

    # Feladat
    add_heading(doc, "A megvalósítandó feladat", level=2)
    add_para(doc, cfg["project_title"], bold=True)

    add_para(doc, "")

    add_para(doc, "A Vállalkozó az alábbi feladatokat valósítja meg:")
    for i, task in enumerate(cfg["tasks"], 1):
        add_numbered(doc, task)

    add_para(doc, "")

    # Tech stack
    if cfg["tech_stack"]:
        p = doc.add_paragraph()
        p.add_run("Technológia: ").bold = True
        p.add_run(cfg["tech_stack"])
        add_para(doc, "")

    # Üzemeltetési költségek
    if cfg["third_party_costs"]:
        add_heading(doc, "Üzemeltetési költségek (harmadik fél, tájékoztató jellegű)", level=2)
        for cost in cfg["third_party_costs"]:
            add_bullet(doc, f"{cost['item']}: {cost['estimate']}")
        add_para(doc,
                 "A fenti költségek a Megrendelőt terhelik, és a harmadik fél mindenkor hatályos árazása szerint alakulnak. A Vállalkozói Díj ezeket nem tartalmazza.",
                 italic=True)
        add_para(doc, "")

    # Vállalkozói díj
    add_heading(doc, "Vállalkozói Díj", level=2)
    p = doc.add_paragraph()
    p.add_run(f"{cfg['fee_net']} + ÁFA").bold = True
    if cfg.get("fee_text"):
        p.add_run(f" (azaz {cfg['fee_text']} forint + ÁFA)")

    add_para(doc, "")

    # Fizetési ütemezés
    add_heading(doc, "Fizetési ütemezés", level=2)
    add_para(doc, cfg["payment_schedule"])

    add_para(doc, "")

    # Határidő
    if cfg["deadline"]:
        add_heading(doc, "Teljesítési határidő", level=2)
        add_para(doc, cfg["deadline"])
        add_para(doc, "")

    # Egyéb rendelkezés (boilerplate)
    add_heading(doc, "Egyéb rendelkezések", level=2)
    add_para(doc,
             "A jelen Munkalapra a Keretszerződés rendelkezései irányadók. A teljesítés elfogadása a Keretszerződés 4. pontja szerinti Teljesítési Igazolás aláírásával történik. A forráskód és a létrehozott anyagok a Vállalkozói Díj megfizetésével egyidejűleg kerülnek véglegesen átadásra a Megrendelő részére.")

    add_para(doc, "")
    add_para(doc, "")

    # Aláírás
    if cfg.get("year") and cfg.get("month"):
        add_para(doc, f"{cfg['city']}, {cfg['year']}. {cfg['month']} ...")
    else:
        add_para(doc, f"{cfg['city']}, ............................")

    add_para(doc, "")
    add_para(doc, "")

    # Két aláírás oszlop helyett egymás után, mert sima docx egyszerűbb
    add_signature_line(doc, "Megrendelő", cfg["company_name"], "")
    add_para(doc, "")
    add_signature_line(doc, "Vállalkozó", "BILDR HUB Kft.", "képv.: Kovács Bence, ügyvezető")

    doc.save(str(output_path))
    print(f"OK: {output_path}")


if __name__ == "__main__":
    main()
